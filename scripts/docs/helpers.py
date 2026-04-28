# -*- coding: utf-8 -*-
"""
Shared helper module for AgriConnect thesis document generation.
All formatting, font, spacing, and numPr fixes in one place.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = "UT Sans"
OUTPUT = r"c:\Users\Admin\Desktop\Licenta Adriana\scripts\docs\AgriConnect_Diploma_Final.docx"


def open_doc():
    """Open the document and re-apply numPr removal + line spacing defaults."""
    doc = Document(OUTPUT)
    _remove_numpr(doc)
    return doc


def _remove_numpr(doc):
    """Remove auto-numbering from all Heading styles to prevent '4 1.1 INTRODUCERE'."""
    styles_part = doc.styles.element
    for style_el in styles_part.iter(qn('w:style')):
        style_id = style_el.get(qn('w:styleId'), '')
        if style_id.startswith('Heading') or style_id.startswith('heading'):
            pPr = style_el.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)


def set_font(run, name=FONT_NAME, size=12, bold=False, italic=False, color=None):
    """Apply UT Sans font formatting to a run with XML-level enforcement."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" '
            f'w:cs="{name}" w:eastAsia="{name}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        for attr in ['w:ascii', 'w:hAnsi', 'w:cs']:
            rFonts.set(qn(attr), name)


def _apply_line_spacing(p, spacing=1.5):
    """Set 1.5 line spacing on a paragraph."""
    pf = p.paragraph_format
    pf.line_spacing = spacing


def h1(doc, text):
    """Add Heading 1 with UT Sans, 18pt, 1.5 line spacing."""
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in h.runs:
        set_font(r, size=18, bold=True)
    pf = h.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_together = True
    pf.keep_with_next = True
    _apply_line_spacing(h)
    return h


def h2(doc, text):
    """Add Heading 2 with UT Sans, 14pt, 1.5 line spacing."""
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in h.runs:
        set_font(r, size=14, bold=True)
    pf = h.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(0)
    pf.keep_together = True
    pf.keep_with_next = True
    _apply_line_spacing(h)
    return h


def h3(doc, text):
    """Add Heading 3 with UT Sans, 11pt bold, 1.5 line spacing."""
    h = doc.add_heading(text, level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in h.runs:
        set_font(r, size=11, bold=True)
    pf = h.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(0)
    pf.keep_together = True
    pf.keep_with_next = True
    _apply_line_spacing(h)
    return h


def para(doc, text, size=12, bold=False, italic=False, sa=None):
    """Add a justified paragraph with UT Sans, 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    _apply_line_spacing(p)
    if sa is not None:
        p.paragraph_format.space_after = Pt(sa)
    return p


def mixed(doc, segments):
    """Add a paragraph with mixed formatting segments: [(text, bold, italic, size), ...]."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in segments:
        t = seg[0]
        b = seg[1] if len(seg) > 1 else False
        it = seg[2] if len(seg) > 2 else False
        sz = seg[3] if len(seg) > 3 else 12
        run = p.add_run(t)
        set_font(run, size=sz, bold=b, italic=it)
    _apply_line_spacing(p)
    return p


def bullet(doc, text, size=11):
    """Add a manual bullet point with indent and 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
    pPr.append(ind)
    run = p.add_run("\u2022 " + text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(4)
    _apply_line_spacing(p)
    return p


def page_break(doc):
    """Insert a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_table(doc, headers, rows):
    """Add a formatted table with UT Sans font."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for r in hdr[i].paragraphs[0].runs:
            set_font(r, size=11, bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for p in row[i].paragraphs:
                for r in p.runs:
                    set_font(r, size=11)
    para(doc, "", sa=4)
    return table


def bib_entry(doc, num, text):
    """Add a bibliography entry with hanging indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Hanging indent: 1cm left, 1cm hanging
    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="567" w:hanging="567"/>')
    pPr.append(ind)
    run = p.add_run(f"[{num}] {text}")
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    _apply_line_spacing(p)
    return p


def save_doc(doc):
    """Save the document."""
    doc.save(OUTPUT)

def add_code_block(doc, code_str, lang="typescript"):
    """Add a block of code with Courier New font, single spacing, and light gray background."""
    from docx.shared import RGBColor
    
    # Add an empty paragraph before for spacing
    p_before = doc.add_paragraph()
    p_before.paragraph_format.space_after = Pt(2)
    
    for line in code_str.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0 # Single space for code
        p.paragraph_format.left_indent = Pt(20)
        
        # We can simulate syntax highlighting by using different colors, 
        # but for simplicity we will just make it all Courier New.
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        
        # Force font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New" '
                f'w:cs="Courier New" w:eastAsia="Courier New"/>'
            )
            rPr.insert(0, rFonts)
        else:
            for attr in ['w:ascii', 'w:hAnsi', 'w:cs']:
                rFonts.set(qn(attr), "Courier New")
                
    # Add an empty paragraph after
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)

