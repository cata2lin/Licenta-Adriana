# -*- coding: utf-8 -*-
"""
Step 0: Prepare the clean base document from the university template.
- Opens official template
- Preserves cover pages (elements 0..CUT_AFTER)
- Removes all guide/demo content
- Removes numPr from heading styles
- Adds front matter (CUPRINS placeholder, Lista Acronime)
- Saves as AgriConnect_Diploma_Final.docx
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEMPLATE = r"c:\Users\Admin\Desktop\Licenta Adriana\documentatie\Proiect_Diploma_ghid_2025-2026_RO.docx"
OUTPUT = r"c:\Users\Admin\Desktop\Licenta Adriana\scripts\docs\AgriConnect_Diploma_Final.docx"
FONT_NAME = "UT Sans"

print("=" * 60)
print("STEP 0: Preparing clean base from university template")
print("=" * 60)

# --- 1. Open template ---
doc = Document(TEMPLATE)
body = doc.element.body
all_elements = list(body)
print(f"  Template loaded: {len(all_elements)} body elements")

# --- 2. Find cover page section break ---
sect_in_para = []
for i, el in enumerate(all_elements):
    if el.tag.endswith('}p'):
        pPr = el.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            sect_in_para.append(i)

CUT_AFTER = sect_in_para[0] if sect_in_para else 76
print(f"  Cover page ends at element index: {CUT_AFTER}")

# --- 3. Preserve final sectPr (document settings) ---
final_sectPr = None
for el in reversed(all_elements):
    if el.tag == qn('w:sectPr'):
        final_sectPr = el
        break

# --- 4. Remove all guide content after covers ---
removed = 0
for el in all_elements[CUT_AFTER + 1:]:
    if el == final_sectPr:
        continue
    body.remove(el)
    removed += 1
print(f"  Removed {removed} guide elements, kept {CUT_AFTER + 1} cover elements")

# --- 5. Remove numPr from ALL heading styles ---
styles_part = doc.styles.element
numpr_removed = 0
for style_el in styles_part.iter(qn('w:style')):
    style_id = style_el.get(qn('w:styleId'), '')
    if style_id.startswith('Heading') or style_id.startswith('heading'):
        pPr = style_el.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
                numpr_removed += 1
                print(f"    Removed numPr from style '{style_id}'")
print(f"  Total numPr removals: {numpr_removed}")

# --- Helper for front matter ---
def sf(run, size=12, bold=False, italic=False):
    run.font.name = FONT_NAME; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        for a in ['w:ascii','w:hAnsi','w:cs']: rFonts.set(qn(a), FONT_NAME)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sz = {1:18, 2:14, 3:11}.get(level, 12)
    for r in h.runs: sf(r, size=sz, bold=True)
    pf = h.paragraph_format
    pf.space_before = Pt(18); pf.space_after = Pt(8)
    pf.keep_together = True; pf.keep_with_next = True
    pf.line_spacing = 1.5

def add_p(text, size=12, bold=False, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text); sf(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.line_spacing = 1.5

def pb():
    p = doc.add_paragraph(); run = p.add_run(); run.add_break(WD_BREAK.PAGE)

# --- 6. Front Matter ---
print("  Adding front matter...")

# CUPRINS
pb()
add_h("CUPRINS", level=1)
add_p("[Cuprinsul se va genera automat din Word: References → Table of Contents → Update Table]", size=11, italic=True)
pb()

# LISTA DE FIGURI
add_h("LISTA DE FIGURI, TABELE ȘI CODURI SURSĂ", level=1)
add_p("[Listele se vor genera automat din Word: References → Insert Table of Figures]", size=11, italic=True)
pb()

# LISTA DE ACRONIME
add_h("LISTA DE ACRONIME", level=1)
acronyms = [
    ("ADR", "Alternative Dispute Resolution — Rezolvare Alternativă a Disputelor"),
    ("ANAF", "Agenția Națională de Administrare Fiscală"),
    ("ANSVSA", "Autoritatea Națională Sanitară Veterinară și pentru Siguranța Alimentelor"),
    ("API", "Application Programming Interface"),
    ("BRFR", "Bunuri cu Risc Fiscal Ridicat"),
    ("CAEN", "Clasificarea Activităților din Economia Națională"),
    ("CAP/PAC", "Common Agricultural Policy / Politica Agricolă Comună"),
    ("CH", "Contraction Hierarchies"),
    ("CUI", "Cod Unic de Identificare"),
    ("DAC7", "Directiva (UE) 2021/514 privind cooperarea administrativă în domeniul fiscal"),
    ("DDD", "Domain-Driven Design"),
    ("DI", "Dependency Injection"),
    ("DTO", "Data Transfer Object"),
    ("EUDI", "European Digital Identity"),
    ("GDPR", "General Data Protection Regulation"),
    ("GIN", "Generalized Inverted Index"),
    ("IAM", "Identity and Access Management"),
    ("JSONB", "JavaScript Object Notation — Binary"),
    ("JWT", "JSON Web Token"),
    ("KYC/KYB", "Know Your Customer / Know Your Business"),
    ("MATIF", "Marché à Terme International de France"),
    ("NC", "Nomenclatura Combinată"),
    ("OSRM", "Open Source Routing Machine"),
    ("PSD2", "Payment Services Directive 2"),
    ("RBAC", "Role-Based Access Control"),
    ("RO e-Transport", "Sistemul național privind monitorizarea transporturilor rutiere de bunuri"),
    ("RSC", "React Server Components"),
    ("SCA", "Strong Customer Authentication"),
    ("SPA", "Single Page Application"),
    ("SWR", "Stale-While-Revalidate"),
    ("TOAST", "The Oversize Attribute Storage Technique"),
    ("TSP", "Traveling Salesperson Problem"),
    ("TVA", "Taxa pe Valoarea Adăugată"),
    ("UIT", "Codul Unic de Identificare a Transportului"),
    ("UTP", "Unfair Trading Practices — Practici Comerciale Neloiale"),
]

for acr, full in acronyms:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_acr = p.add_run(f"{acr}")
    sf(run_acr, size=11, bold=True)
    run_sep = p.add_run(f" — {full};")
    sf(run_sep, size=11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5

pb()

# --- 7. Save clean base ---
doc.save(OUTPUT)
print(f"\n✅ Clean base saved: {OUTPUT}")

# Verify
doc2 = Document(OUTPUT)
print(f"   Paragraphs: {len(doc2.paragraphs)}")
print(f"   Elements: {len(list(doc2.element.body))}")
print("   Ready for chapter generation scripts.")
