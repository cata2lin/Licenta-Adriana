# -*- coding: utf-8 -*-
"""
Master Runner: Executes the full thesis generation pipeline in order.
1. rebuild_base.py  — Clean template, front matter
2. gen_ch1.py       — Chapter 1: Introducere
3. gen_ch2.py       — Chapter 2: Arhitectura
4. gen_ch3.py       — Chapter 3: Cadrul Legislativ
5. gen_ch4.py       — Chapter 4: Implementarea
6. gen_ch5_bib.py   — Chapter 5 + Bibliografie
"""
import subprocess, sys, os, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SCRIPTS = [
    "rebuild_base.py",
    "gen_ch1.py",
    "gen_ch2.py",
    "gen_ch3.py",
    "gen_ch4.py",
    "gen_ch5_bib.py",
]

cwd = os.path.dirname(os.path.abspath(__file__))
print("=" * 60)
print("  AgriConnect Thesis — Full Pipeline Runner")
print("=" * 60)

for i, script in enumerate(SCRIPTS, 1):
    print(f"\n{'─' * 50}")
    print(f"  [{i}/{len(SCRIPTS)}] Running: {script}")
    print(f"{'─' * 50}")
    result = subprocess.run(
        [sys.executable, script],
        cwd=cwd,
        capture_output=True,
        text=True,
        encoding='utf-8'
    )
    if result.stdout:
        print(result.stdout.strip())
    if result.returncode != 0:
        print(f"\n❌ FAILED: {script}")
        if result.stderr:
            print(result.stderr)
        sys.exit(1)

print(f"\n{'═' * 60}")
print("  ✅ Pipeline complete! All chapters generated successfully.")
print(f"{'═' * 60}")

# Final stats
from docx import Document
doc = Document(os.path.join(cwd, "AgriConnect_Diploma_Final.docx"))
paras = len(doc.paragraphs)
headings = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
tables = len(doc.tables)
chars = sum(len(p.text) for p in doc.paragraphs)
print(f"\n  Paragraphs: {paras}")
print(f"  Headings:   {headings}")
print(f"  Tables:     {tables}")
print(f"  Characters: {chars:,}")
print(f"  Est. Pages: {chars / 2500:.0f}+ (with 1.5 line spacing, tables, margins)")
