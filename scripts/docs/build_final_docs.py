# -*- coding: utf-8 -*-
"""
Full Document Generator - AgriConnect B2B Marketplace
Folosește template-ul oficial, curăță conținutul demo, elimină auto-numerotarea (pt a nu avea "4 1. Introducere") 
și adaugă manual tot conținutul capitolelor 1-4.
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEMPLATE = r"c:\Users\Admin\Desktop\Licenta Adriana\documentatie\Proiect_Diploma_ghid_2025-2026_RO.docx"
OUTPUT = r"c:\Users\Admin\Desktop\Licenta Adriana\scripts\docs\AgriConnect_Diploma_Final.docx"
FONT_NAME = "UT Sans"

print("1. Inceperea generarii...")

doc = Document(TEMPLATE)
body = doc.element.body
all_elements = list(body)

# --- 1. Clean Template Content ---
sect_in_para = []
for i, el in enumerate(all_elements):
    if el.tag.endswith('}p'):
        pPr = el.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            sect_in_para.append(i)

CUT_AFTER = sect_in_para[0] if sect_in_para else 76

final_sectPr = None
for el in reversed(all_elements):
    if el.tag == qn('w:sectPr'):
        final_sectPr = el
        break

removed = 0
for el in all_elements[CUT_AFTER + 1:]:
    if el == final_sectPr:
        continue
    body.remove(el)
    removed += 1

print(f"Template curățat: Am pastrat {CUT_AFTER + 1} elemente (copertile), am eliminat {removed} elemente vechi.")

# --- 2. Remove Auto-Numbering from Headings ---
styles_part = doc.styles.element
for style_el in styles_part.iter(qn('w:style')):
    style_id = style_el.get(qn('w:styleId'), '')
    if style_id.startswith('Heading') or style_id.startswith('heading'):
        pPr = style_el.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)

# --- Helpers ---
def set_font(run, name=FONT_NAME, size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}" w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in h.runs:
        set_font(run, size={1:18, 2:14, 3:11}.get(level, 12), bold=True)
    pf = h.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
    elif level == 2:
        pf.space_before = Pt(18)
        pf.space_after = Pt(0)
    elif level == 3:
        pf.space_before = Pt(10)
        pf.space_after = Pt(0)
    pf.keep_together = True
    pf.keep_with_next = True
    return h

def add_para(doc, text, size=12, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p

# --- 3. Front Matter ---
add_page_break(doc)
add_heading_styled(doc, "CUPRINS", level=1)
add_para(doc, "[Cuprinsul se va genera automat din Word: References -> Table of Contents -> Update Table]", size=11, italic=True)
add_page_break(doc)

add_heading_styled(doc, "LISTA DE FIGURI, TABELE ȘI CODURI SURSĂ", level=1)
add_para(doc, "[Listele se vor genera automat din Word: References -> Insert Table of Figures]", size=11, italic=True)
add_page_break(doc)

add_heading_styled(doc, "LISTA DE ACRONIME", level=1)
acronyms = [
    ("ANAF", "Agenția Națională de Administrare Fiscală"),
    ("API", "Application Programming Interface"),
    ("DDD", "Domain-Driven Design"),
    ("DAC7", "Directiva (UE) 2021/514 privind cooperarea administrativă în domeniul fiscal"),
    ("GDPR", "General Data Protection Regulation"),
    ("IAM", "Identity and Access Management"),
    ("JWT", "JSON Web Token"),
    ("OSRM", "Open Source Routing Machine"),
    ("PSD2", "Payment Services Directive 2"),
    ("RBAC", "Role-Based Access Control"),
    ("RO e-Transport", "Sistemul național privind monitorizarea transporturilor rutiere de bunuri cu risc fiscal ridicat"),
    ("UIT", "Codul Unic de Identificare a Transportului"),
]

for acr, full in acronyms:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_acr = p.add_run(f"{acr}")
    set_font(run_acr, size=11, bold=True)
    run_sep = p.add_run(f" – {full};")
    set_font(run_sep, size=11)
    p.paragraph_format.space_after = Pt(2)

add_page_break(doc)

# --- 4. Chapter Content ---
print("Adaugare Capitolul 1...")
add_heading_styled(doc, '1. INTRODUCERE ȘI STADIUL ACTUAL AL PIEȚEI', 1)
add_heading_styled(doc, '1.1 Contextul Global și Lanțul Agroalimentar', 2)
add_para(doc, 'Lanțul agroalimentar global se află într-un punct de inflexiune tehnologică. Transformarea digitală a comerțului cu mărfuri agricole, în special pe piața cerealelor din România, reprezintă o evoluție semnificativă de la metodele tradiționale și fragmentate la ecosisteme integrate și bazate pe date. România ocupă o poziție unică în acest peisaj, caracterizându-se prin cel mai mare număr de ferme din Uniunea Europeană — aproximativ 2.9 milioane — și o structură duală ce cuprinde un număr vast de ferme de subzistență alături de un grup restrâns de corporații agroindustriale profesionale care gestionează mai mult de jumătate din terenul agricol național.')
add_para(doc, 'Istoric, tranzacționarea a fost dominată de piețele fizice (cash markets) și piețele futures, precum Marché à Terme International de France (MATIF), care servește drept benchmark european pentru prețul cerealelor. În contracte lichide precum "Euronext Milling Wheat No. 2", mecanismul livrării fizice (market of last resort) asigură ancorarea instrumentului financiar la marfa subiacentă. Totuși, decuplarea dintre semnalele financiare și fluxul fizic creează volatilitate, iar lanțul tradițional de aprovizionare rămâne fragmentat, implicând adesea până la șase intermediari între fermier și procesatorul final.')

add_heading_styled(doc, '1.2 Disintermedierea și Nevoia unei Platforme B2B', 2)
add_para(doc, 'Problemele fundamentale ale lanțului clasic includ opacitatea prețurilor, costurile logistice ridicate și riscul de contraparte semnificativ. Intermediarii adaugă margini comerciale fără a adăuga valoare proporțională în termeni de calitate sau trasabilitate, erodând profitul producătorului primar. Digitalizarea prin marketplace-uri B2B promite disintermedierea directă, permițând recuperarea unor margini de profit cuprinse între 10% și 25%. Spre deosebire de comerțul electronic B2C, platformele AgTech B2B trebuie să gestioneze tranzacții de mare valoare (depășind adesea 1 milion USD pentru cargouri vrac), impunând mecanisme riguroase de securitate, escrow și conformitate legală (NACE 4611).')

add_heading_styled(doc, '1.3 Standarde de Calitate a Cerealelor: SR EN 15587', 2)
add_para(doc, 'Fungibilitatea mărfurilor agricole pe bursele internaționale depinde de standardizarea riguroasă a parametrilor de calitate. În Europa, principalul benchmark pentru puritatea cerealelor este standardul EN 15587, care definește termenul "Besatz" (impurități). Acesta cuprinde toate materiile străine dintr-o mostră, incluzând boabe sparte, zbârcite, afectate de dăunători, semințe de buruieni ("Schwarzbesatz") și fragmente minerale.')
add_para(doc, 'Standardul impune un proces riguros de sitare (sieve analysis) utilizând site cu orificii de 2.00 mm pentru grâul comun. În România, piața este profund influențată de agrometeorologie; de exemplu, deși anul agricol 2024 a fost afectat de secetă extremă, calitatea cerealelor (conținutul de proteină măsurat prin metoda Kjeldahl și glutenul) a fost superioară din cauza stresului hidric, în timp ce producția din 2025 a înregistrat o recuperare spectaculoasă a volumelor.')

add_page_break(doc)

print("Adaugare Capitolul 2...")
add_heading_styled(doc, '2. ARHITECTURA SISTEMULUI ȘI TEHNOLOGII', 1)
add_heading_styled(doc, '2.1 Strategic Design: Domain-Driven Design (DDD)', 2)
add_para(doc, 'Integritatea arhitecturală a unui marketplace multi-sided B2B depinde de capacitatea sa de a gestiona logica complexă fără a acumula datorie tehnică. Pentru AgriConnect, s-a adoptat Domain-Driven Design (DDD), separând sistemul în mai multe "Bounded Contexts" izolate: Identity and Access Management (IAM), Trading, Financial, Logistics, și Dispute Resolution. Fiecare context are propriul său vocabular ubicu ("Ubiquitous Language"), asigurând că termeni precumi "Contract" sau "Besatz" își păstrează sensul tehnic și comercial nealterat.')
add_para(doc, 'Sistemul este construit ca un "Modular Monolith". Spre deosebire de abordarea pură bazată pe microservicii, monolitul modular păstrează modulele în aceeași unitate de execuție, dar izolate logic (Hexagonal Architecture / Ports and Adapters). Comunicarea între domenii se realizează prin Domain Events și Integration Events folosind Outbox Pattern, evitând problemele sistemelor distribuite.')

add_heading_styled(doc, '2.2 Backend: NestJS și Fastify', 2)
add_para(doc, 'Alegerea NestJS cu TypeScript permite transpunerea tactică a DDD prin utilizarea agregatelor, entităților și a obiectelor valorice (Value Objects). Sistemul de Dependency Injection (DI) din NestJS izolează logica de business în clase de tip Service și controlează strict vizibilitatea publică a modulelor. Pentru a atinge performanțe enterprise, a fost integrat adaptorul HTTP Fastify, care oferă un throughput de procesare a request-urilor de până la de 3 ori mai mare comparativ cu soluția implicită Express.js.')

add_heading_styled(doc, '2.3 Baze de Date: Stocare Hibridă cu PostgreSQL JSONB', 2)
add_para(doc, 'Diversitatea parametrilor calitativi agricoli impune o schemă hibridă. PostgreSQL a fost ales pentru garanțiile sale ACID tranzacționale. Datele statice (ID-uri, asocieri, preț de bază) sunt stocate în coloane relaționale, în timp ce atributele dinamice specifice mărfii (umiditate, proteină, Besatz) sunt stocate în câmpuri JSONB (Binary JSON).')
add_para(doc, 'Pentru performanța interogărilor, platforma utilizează indecși GIN (Generalized Inverted Index) cu operatorul jsonb_path_ops, permițând căutări de tip "containment" (ex: filtrarea listărilor care au proteina peste 12.5%) extrem de rapide. O atenție deosebită este acordată tehnicii TOAST (The Oversize Attribute Storage Technique) din Postgres și funcțiilor native IS JSON introduse în versiunile recente pentru a asigura integritatea datelor ("JSON Schema Validation").')

add_heading_styled(doc, '2.4 Frontend: React 19, Server Components și UI Premium', 2)
add_para(doc, 'Arhitectura frontend-ului folosește React 19 și noile paradigme React Server Components (RSC) pentru a muta execuția și orchestarea stării de la client pe server. Gestiunea stării globale dinamice este delegată bibliotecii Zustand. Designul vizual este concentrat pe "Dark Mode" (o preferință pentru 82% din utilizatorii B2B enterprise pentru reducerea oboselii oculare) și "Glassmorphism", utilizând backdrop-filter pentru a crea ierarhii vizuale distincte și profunzime fără a compromite lizibilitatea. Pentru vizualizarea datelor financiare complexe, se utilizează diagrame bazate pe SVG (Scalable Vector Graphics), asigurând rendering clar la orice rezoluție.')

add_heading_styled(doc, '2.5 Securitate și Autentificare (eIDAS 2.0 și Argon2id)', 2)
add_para(doc, 'Platforma implementează autentificare stateless via JWT (JSON Web Tokens). Pentru rezolvarea limitărilor privind revocarea token-urilor, este utilizată o arhitectură "Theme Park Ticket": access tokens cu durată scurtă (5-15 min) și refresh tokens rotative cu invalidare într-un blacklist pe Redis.')
add_para(doc, 'Protejarea parolelor este realizată prin algoritmul "Memory-Hard" Argon2id (câștigătorul Password Hashing Competition), esențial într-o eră în care capabilitățile de brute-forcing ale plăcilor video (ex: clustere de RTX 5090) fac algoritmii vechi precum bcrypt mai puțin siguri. Mai mult, arhitectura platformei prefigurează compatibilitatea cu eIDAS 2.0 și protocolul OpenID4VP, care din 2027 va impune acceptarea European Digital Identity (EUDI) Wallet și a credențialelor verificabile bazate pe "Selective Disclosure".')

add_page_break(doc)

print("Adaugare Capitolul 3...")
add_heading_styled(doc, '3. CADRUL LEGISLATIV, FISCAL ȘI CALITATIV', 1)

add_heading_styled(doc, '3.1 Regimul Fiscal: Taxarea Inversă (Art. 331)', 2)
add_para(doc, 'Comerțul cu cereale în România este supus unui regim fiscal special definit de Articolul 331 din Codul Fiscal: taxarea inversă. Această măsură antifraudă mută responsabilitatea plății TVA de la furnizor la beneficiar pentru anumite coduri din Nomenclatura Combinată (NC), cum ar fi grâul (1001), porumbul (1005) sau floarea-soarelui (1206). AgriConnect automatizează acest mecanism verificând statusul TVA al participanților și emițând facturile cu mențiunea obligatorie, având un efect neutru asupra trezoreriei.')

add_heading_styled(doc, '3.2 Transparență și Monitorizare: DAC7 și RO e-Transport', 2)
add_para(doc, 'Sub incidența Directivei UE 2021/514 (DAC7) și OG 16/2023, platformele digitale au obligația de a raporta anual veniturile vânzătorilor la ANAF. Platforma colectează și raportează automat datele de identificare și volumul tranzacțiilor financiare ale fermierilor (formular F7000).')
add_para(doc, 'Sistemul național RO e-Transport monitorizează transporturile rutiere cu risc fiscal ridicat (BRFR). Pentru vehiculele peste 2.5 tone, cu marfă peste 500 kg sau valoare > 10.000 lei, sistemul generează automat un cod UIT (Unic de Identificare a Transportului) prin integrarea API-ului ANAF v2 în modulul de logistică, transferând inclusiv datele GPS de poziționare conform reglementărilor din 2025.')

add_heading_styled(doc, '3.3 Infrastructura Financiară: PSD2 și Sistemul Escrow', 2)
add_para(doc, 'Reglementată de Directiva PSD2, platforma funcționează cu conturi de segregare (Safeguarded Accounts). Fondurile cumpărătorului sunt reținute într-un cont neutru de către furnizorii de servicii de plată (TPP), izolate de activele marketplace-ului. Banii sunt eliberați vânzătorului doar după confirmarea digitală a recepției logistice. Acest sistem Escrow digital elimină riscul de contraparte, validând fiecare eliberare printr-o Autentificare Strictă a Clienților (SCA).')

add_heading_styled(doc, '3.4 Practici Comerciale Neloiale și Delimitarea de Bursele Reglementate', 2)
add_para(doc, 'Platforma B2B respectă prevederile Legii 81/2022 (transpunerea Directivei UE 2019/633), protejând fermierii mici de practicile neloiale (ex: întârzieri la plată, taxe mascate de listare). Mai mult, AgriConnect funcționează ca un intermediar digital (CAEN 4611) pentru contracte spot, delimitându-se juridic de Bursa de Mărfuri (Legea 357/2005) reglementată de ASF, evitând obligativitatea administrării unei piețe de interes public centralizate.')

add_heading_styled(doc, '3.5 GDPR, Privacy by Design și Trasabilitate (Farm-to-Fork)', 2)
add_para(doc, 'Design-ul sistemului respectă conceptul "Privacy by Design", minimizând și pseudonimizând informațiile sensibile stocate. În conformitate cu Regulamentul (CE) 178/2002 privind siguranța alimentară, arhitectura asigură trasabilitatea "one-step-forward, one-step-back", menținând registre inalterabile ale calității și provenienței loturilor de cereale pentru certificare ANSVSA.')

add_page_break(doc)

print("Adaugare Capitolul 4...")
add_heading_styled(doc, '4. IMPLEMENTAREA PRACTICĂ ȘI OPTIMIZAREA LOGISTICĂ', 1)

add_heading_styled(doc, '4.1 Fluxul Practic: Bounded Contexts în Acțiune', 2)
add_para(doc, 'Platforma orchestrează mai multe subdomenii. Modulul de "Trading" gestionează ofertele (RFQ) și contractarea asincronă, în timp ce modulul "IAM" procesează permisiunile de tip RBAC. Odată finalizat contractul, un eveniment este declanșat către "Financial" pentru alocarea Escrow și către "Transport" pentru generarea rutei și a codului UIT.')

add_heading_styled(doc, '4.2 OSRM și Optimizarea Transporturilor (Digital Freight Matching)', 2)
add_para(doc, 'Costul de logistică poate decima profiturile. Pentru calcularea matricelor de distanță și rutare, platforma folosește motorul Open Source Routing Machine (OSRM). Utilizând algoritmul "Contraction Hierarchies" (CH) sau "Multi-Level Dijkstra", graful rețelei stradale (OSM) este preprocesat în RAM, permițând interogări în milisecunde.')
add_para(doc, 'Sistemul de Freight Matching integrează algoritmi de învățare automată (Reinforcement Learning / XGBoost) și algoritmi TSP (Traveling Salesperson) pentru optimizarea ridicării mărfurilor din multiple ferme, reducând "deadhead miles" (kilometraj în gol). Identificarea "last-meter" este gestionată prin Google Places Autocomplete API, eliminând erorile de navigație industrială.')

add_heading_styled(doc, '4.3 Reziliența Platformei și Fallback Mechanisms', 2)
add_para(doc, 'Pentru o stabilitate de 99.9%, arhitectura folosește containerizare Docker și metode Agile Scrum pentru deployment continuu. Pentru a preveni erorile în cascadă atunci când API-urile externe (cum ar fi e-Transport ANAF) suferă degradări de performanță, sistemul implementează "Circuit Breaker Pattern". În interfață, se folosește directiva "Stale-While-Revalidate" (SWR) pentru a afișa instant date din cache în timp ce se revalidează fluxul de prețuri în fundal.')

add_heading_styled(doc, '4.4 Sistemul de Reputație și RXM', 2)
add_para(doc, 'Încrederea este codificată prin algoritmi de Reputation Experience Management (RXM). Evaluările utilizatorilor nu reprezintă o simplă medie, ci aplică algoritmul "Weighted Liquid Rank". Astfel, greutatea recenziei depinde de valoarea tranzacției, degradarea temporală (Time Decay) și scorul de integritate istoric al evaluatorului, creând protecție eficientă contra atacurilor de manipulare (shilling/smearing).')

add_page_break(doc)
add_heading_styled(doc, 'BIBLIOGRAFIE', 1)
add_para(doc, '[Bibliografie - Va fi adaugata automat din sistemul de citari]')

doc.save(OUTPUT)
print(f"Generare cu succes! Document salvat la {OUTPUT}")
