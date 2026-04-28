# -*- coding: utf-8 -*-
"""Quick check: count paragraphs, headings, tables in the final doc."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r"c:\Users\Admin\Desktop\Licenta Adriana\scripts\docs\AgriConnect_Diploma_Final.docx")

paras = len(doc.paragraphs)
headings = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
tables = len(doc.tables)
chars = sum(len(p.text) for p in doc.paragraphs)

# Estimate pages: ~3000 chars per page (12pt, margins)
est_pages = chars / 3000

print(f"Paragraphs: {paras}")
print(f"Headings:   {headings}")
print(f"Tables:     {tables}")
print(f"Characters: {chars:,}")
print(f"Est. Pages: {est_pages:.0f}")
