import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def main():
    file_path = 'AgriConnect_Diploma_v1.docx'
    if not os.path.exists(file_path):
        print(f"Eroare: Nu am gasit {file_path}")
        return

    doc = Document(file_path)

    doc.add_page_break()

    # CAPITOLUL 1
    add_heading(doc, 'Capitolul 1: Introducere și Stadiul Actual al Pieței (Literature Review)', 1)
    
    add_heading(doc, '1.1 Contextul Global și Lanțul Agroalimentar', 2)
    add_paragraph(doc, 'Lanțul agroalimentar global se află într-un punct de inflexiune tehnologică. Transformarea digitală a comerțului cu mărfuri agricole, în special pe piața cerealelor din România, reprezintă o evoluție semnificativă de la metodele tradiționale și fragmentate la ecosisteme integrate și bazate pe date. România ocupă o poziție unică în acest peisaj, caracterizându-se prin cel mai mare număr de ferme din Uniunea Europeană — aproximativ 2.9 milioane — și o structură duală ce cuprinde un număr vast de ferme de subzistență alături de un grup restrâns de corporații agroindustriale profesionale care gestionează mai mult de jumătate din terenul agricol național.')
    add_paragraph(doc, 'Istoric, tranzacționarea a fost dominată de piețele fizice (cash markets) și piețele futures, precum Marché à Terme International de France (MATIF), care servește drept benchmark european pentru prețul cerealelor. În contracte lichide precum "Euronext Milling Wheat No. 2", mecanismul livrării fizice (market of last resort) asigură ancorarea instrumentului financiar la marfa subiacentă. Totuși, decuplarea dintre semnalele financiare și fluxul fizic creează volatilitate, iar lanțul tradițional de aprovizionare rămâne fragmentat, implicând adesea până la șase intermediari între fermier și procesatorul final.')
    
    add_heading(doc, '1.2 Disintermedierea și Nevoia unei Platforme B2B', 2)
    add_paragraph(doc, 'Problemele fundamentale ale lanțului clasic includ opacitatea prețurilor, costurile logistice ridicate și riscul de contraparte semnificativ. Intermediarii adaugă margini comerciale fără a adăuga valoare proporțională în termeni de calitate sau trasabilitate, erodând profitul producătorului primar. Digitalizarea prin marketplace-uri B2B promite disintermedierea directă, permițând recuperarea unor margini de profit cuprinse între 10% și 25%. Spre deosebire de comerțul electronic B2C, platformele AgTech B2B trebuie să gestioneze tranzacții de mare valoare (depășind adesea 1 milion USD pentru cargouri vrac), impunând mecanisme riguroase de securitate, escrow și conformitate legală (NACE 4611).')

    add_heading(doc, '1.3 Standarde de Calitate a Cerealelor: SR EN 15587', 2)
    add_paragraph(doc, 'Fungibilitatea mărfurilor agricole pe bursele internaționale depinde de standardizarea riguroasă a parametrilor de calitate. În Europa, principalul benchmark pentru puritatea cerealelor este standardul EN 15587, care definește termenul "Besatz" (impurități). Acesta cuprinde toate materiile străine dintr-o mostră, incluzând boabe sparte, zbârcite, afectate de dăunători, semințe de buruieni ("Schwarzbesatz") și fragmente minerale.')
    add_paragraph(doc, 'Standardul impune un proces riguros de sitare (sieve analysis) utilizând site cu orificii de 2.00 mm pentru grâul comun. În România, piața este profund influențată de agrometeorologie; de exemplu, deși anul agricol 2024 a fost afectat de secetă extremă, calitatea cerealelor (conținutul de proteină măsurat prin metoda Kjeldahl și glutenul) a fost superioară din cauza stresului hidric, în timp ce producția din 2025 a înregistrat o recuperare spectaculoasă a volumelor.')

    doc.add_page_break()

    # CAPITOLUL 2
    add_heading(doc, 'Capitolul 2: Arhitectura Sistemului și Tehnologii', 1)
    
    add_heading(doc, '2.1 Strategic Design: Domain-Driven Design (DDD)', 2)
    add_paragraph(doc, 'Integritatea arhitecturală a unui marketplace multi-sided B2B depinde de capacitatea sa de a gestiona logica complexă fără a acumula datorie tehnică. Pentru AgriConnect, s-a adoptat Domain-Driven Design (DDD), separând sistemul în mai multe "Bounded Contexts" izolate: Identity and Access Management (IAM), Trading, Financial, Logistics, și Dispute Resolution. Fiecare context are propriul său vocabular ubicu ("Ubiquitous Language"), asigurând că termeni precumi "Contract" sau "Besatz" își păstrează sensul tehnic și comercial nealterat.')
    add_paragraph(doc, 'Sistemul este construit ca un "Modular Monolith". Spre deosebire de abordarea pură bazată pe microservicii, monolitul modular păstrează modulele în aceeași unitate de execuție, dar izolate logic (Hexagonal Architecture / Ports and Adapters). Comunicarea între domenii se realizează prin Domain Events și Integration Events folosind Outbox Pattern, evitând problemele sistemelor distribuite.')

    add_heading(doc, '2.2 Backend: NestJS și Fastify', 2)
    add_paragraph(doc, 'Alegerea NestJS cu TypeScript permite transpunerea tactică a DDD prin utilizarea agregatelor, entităților și a obiectelor valorice (Value Objects). Sistemul de Dependency Injection (DI) din NestJS izolează logica de business în clase de tip Service și controlează strict vizibilitatea publică a modulelor. Pentru a atinge performanțe enterprise, a fost integrat adaptorul HTTP Fastify, care oferă un throughput de procesare a request-urilor de până la de 3 ori mai mare comparativ cu soluția implicită Express.js.')

    add_heading(doc, '2.3 Baze de Date: Stocare Hibridă cu PostgreSQL JSONB', 2)
    add_paragraph(doc, 'Diversitatea parametrilor calitativi agricoli impune o schemă hibridă. PostgreSQL a fost ales pentru garanțiile sale ACID tranzacționale. Datele statice (ID-uri, asocieri, preț de bază) sunt stocate în coloane relaționale, în timp ce atributele dinamice specifice mărfii (umiditate, proteină, Besatz) sunt stocate în câmpuri JSONB (Binary JSON).')
    add_paragraph(doc, 'Pentru performanța interogărilor, platforma utilizează indecși GIN (Generalized Inverted Index) cu operatorul jsonb_path_ops, permițând căutări de tip "containment" (ex: filtrarea listărilor care au proteina peste 12.5%) extrem de rapide. O atenție deosebită este acordată tehnicii TOAST (The Oversize Attribute Storage Technique) din Postgres și funcțiilor native IS JSON introduse în versiunile recente pentru a asigura integritatea datelor ("JSON Schema Validation").')

    add_heading(doc, '2.4 Frontend: React 19, Server Components și UI Premium', 2)
    add_paragraph(doc, 'Arhitectura frontend-ului folosește React 19 și noile paradigme React Server Components (RSC) pentru a muta execuția și orchestarea stării de la client pe server. Gestiunea stării globale dinamice este delegată bibliotecii Zustand. Designul vizual este concentrat pe "Dark Mode" (o preferință pentru 82% din utilizatorii B2B enterprise pentru reducerea oboselii oculare) și "Glassmorphism", utilizând backdrop-filter pentru a crea ierarhii vizuale distincte și profunzime fără a compromite lizibilitatea. Pentru vizualizarea datelor financiare complexe, se utilizează diagrame bazate pe SVG (Scalable Vector Graphics), asigurând rendering clar la orice rezoluție.')

    add_heading(doc, '2.5 Securitate și Autentificare (eIDAS 2.0 și Argon2id)', 2)
    add_paragraph(doc, 'Platforma implementează autentificare stateless via JWT (JSON Web Tokens). Pentru rezolvarea limitărilor privind revocarea token-urilor, este utilizată o arhitectură "Theme Park Ticket": access tokens cu durată scurtă (5-15 min) și refresh tokens rotative cu invalidare într-un blacklist pe Redis.')
    add_paragraph(doc, 'Protejarea parolelor este realizată prin algoritmul "Memory-Hard" Argon2id (câștigătorul Password Hashing Competition), esențial într-o eră în care capabilitățile de brute-forcing ale plăcilor video (ex: clustere de RTX 5090) fac algoritmii vechi precum bcrypt mai puțin siguri. Mai mult, arhitectura platformei prefigurează compatibilitatea cu eIDAS 2.0 și protocolul OpenID4VP, care din 2027 va impune acceptarea European Digital Identity (EUDI) Wallet și a credențialelor verificabile bazate pe "Selective Disclosure".')

    doc.save('AgriConnect_Diploma_v2.docx')
    print("Capitolele 1 si 2 (v2) au fost generate cu succes.")

if __name__ == '__main__':
    main()
