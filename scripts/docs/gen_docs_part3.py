import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def main():
    file_path = 'AgriConnect_Diploma_v2.docx'
    if not os.path.exists(file_path):
        print(f"Eroare: Nu am gasit {file_path}")
        return

    doc = Document(file_path)
    doc.add_page_break()

    # CAPITOLUL 3
    add_heading(doc, 'Capitolul 3: Cadrul Legislativ, Fiscal și Calitativ', 1)

    add_heading(doc, '3.1 Regimul Fiscal: Taxarea Inversă (Art. 331)', 2)
    add_paragraph(doc, 'Comerțul cu cereale în România este supus unui regim fiscal special definit de Articolul 331 din Codul Fiscal: taxarea inversă. Această măsură antifraudă mută responsabilitatea plății TVA de la furnizor la beneficiar pentru anumite coduri din Nomenclatura Combinată (NC), cum ar fi grâul (1001), porumbul (1005) sau floarea-soarelui (1206). AgriConnect automatizează acest mecanism verificând statusul TVA al participanților și emițând facturile cu mențiunea obligatorie, având un efect neutru asupra trezoreriei.')

    add_heading(doc, '3.2 Transparență și Monitorizare: DAC7 și RO e-Transport', 2)
    add_paragraph(doc, 'Sub incidența Directivei UE 2021/514 (DAC7) și OG 16/2023, platformele digitale au obligația de a raporta anual veniturile vânzătorilor la ANAF. Platforma colectează și raportează automat datele de identificare și volumul tranzacțiilor financiare ale fermierilor (formular F7000).')
    add_paragraph(doc, 'Sistemul național RO e-Transport monitorizează transporturile rutiere cu risc fiscal ridicat (BRFR). Pentru vehiculele peste 2.5 tone, cu marfă peste 500 kg sau valoare > 10.000 lei, sistemul generează automat un cod UIT (Unic de Identificare a Transportului) prin integrarea API-ului ANAF v2 în modulul de logistică, transferând inclusiv datele GPS de poziționare conform reglementărilor din 2025.')

    add_heading(doc, '3.3 Infrastructura Financiară: PSD2 și Sistemul Escrow', 2)
    add_paragraph(doc, 'Reglementată de Directiva PSD2, platforma funcționează cu conturi de segregare (Safeguarded Accounts). Fondurile cumpărătorului sunt reținute într-un cont neutru de către furnizorii de servicii de plată (TPP), izolate de activele marketplace-ului. Banii sunt eliberați vânzătorului doar după confirmarea digitală a recepției logistice. Acest sistem Escrow digital elimină riscul de contraparte, validând fiecare eliberare printr-o Autentificare Strictă a Clienților (SCA).')

    add_heading(doc, '3.4 Practici Comerciale Neloiale și Delimitarea de Bursele Reglementate', 2)
    add_paragraph(doc, 'Platforma B2B respectă prevederile Legii 81/2022 (transpunerea Directivei UE 2019/633), protejând fermierii mici de practicile neloiale (ex: întârzieri la plată, taxe mascate de listare). Mai mult, AgriConnect funcționează ca un intermediar digital (CAEN 4611) pentru contracte spot, delimitându-se juridic de Bursa de Mărfuri (Legea 357/2005) reglementată de ASF, evitând obligativitatea administrării unei piețe de interes public centralizate.')

    add_heading(doc, '3.5 GDPR, Privacy by Design și Trasabilitate (Farm-to-Fork)', 2)
    add_paragraph(doc, 'Design-ul sistemului respectă conceptul "Privacy by Design", minimizând și pseudonimizând informațiile sensibile stocate. În conformitate cu Regulamentul (CE) 178/2002 privind siguranța alimentară, arhitectura asigură trasabilitatea "one-step-forward, one-step-back", menținând registre inalterabile ale calității și provenienței loturilor de cereale pentru certificare ANSVSA.')

    doc.add_page_break()

    # CAPITOLUL 4
    add_heading(doc, 'Capitolul 4: Implementarea Practică și Optimizarea Logistică', 1)

    add_heading(doc, '4.1 Fluxul Practic: Bounded Contexts în Acțiune', 2)
    add_paragraph(doc, 'Platforma orchestrează mai multe subdomenii. Modulul de "Trading" gestionează ofertele (RFQ) și contractarea asincronă, în timp ce modulul "IAM" procesează permisiunile de tip RBAC. Odată finalizat contractul, un eveniment este declanșat către "Financial" pentru alocarea Escrow și către "Transport" pentru generarea rutei și a codului UIT.')

    add_heading(doc, '4.2 OSRM și Optimizarea Transporturilor (Digital Freight Matching)', 2)
    add_paragraph(doc, 'Costul de logistică poate decima profiturile. Pentru calcularea matricelor de distanță și rutare, platforma folosește motorul Open Source Routing Machine (OSRM). Utilizând algoritmul "Contraction Hierarchies" (CH) sau "Multi-Level Dijkstra", graful rețelei stradale (OSM) este preprocesat în RAM, permițând interogări în milisecunde.')
    add_paragraph(doc, 'Sistemul de Freight Matching integrează algoritmi de învățare automată (Reinforcement Learning / XGBoost) și algoritmi TSP (Traveling Salesperson) pentru optimizarea ridicării mărfurilor din multiple ferme, reducând "deadhead miles" (kilometraj în gol). Identificarea "last-meter" este gestionată prin Google Places Autocomplete API, eliminând erorile de navigație industrială.')

    add_heading(doc, '4.3 Reziliența Platformei și Fallback Mechanisms', 2)
    add_paragraph(doc, 'Pentru o stabilitate de 99.9%, arhitectura folosește containerizare Docker și metode Agile Scrum pentru deployment continuu. Pentru a preveni erorile în cascadă atunci când API-urile externe (cum ar fi e-Transport ANAF) suferă degradări de performanță, sistemul implementează "Circuit Breaker Pattern". În interfață, se folosește directiva "Stale-While-Revalidate" (SWR) pentru a afișa instant date din cache în timp ce se revalidează fluxul de prețuri în fundal.')

    add_heading(doc, '4.4 Sistemul de Reputație și RXM', 2)
    add_paragraph(doc, 'Încrederea este codificată prin algoritmi de Reputation Experience Management (RXM). Evaluările utilizatorilor nu reprezintă o simplă medie, ci aplică algoritmul "Weighted Liquid Rank". Astfel, greutatea recenziei depinde de valoarea tranzacției, degradarea temporală (Time Decay) și scorul de integritate istoric al evaluatorului, creând protecție eficientă contra atacurilor de manipulare (shilling/smearing).')

    # Placeholder Bibliografie
    doc.add_page_break()
    add_heading(doc, 'Bibliografie', 1)
    p_bib = add_paragraph(doc, '[Aici se va introduce automatizarea bibliografiei sau preluarea referințelor din excel. Citările vor urma standardele academice IEEE sau APA, reflectând legislația din Romania și conceptele DDD, OSRM, JWT].')
    p_bib.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('AgriConnect_Diploma_Final.docx')
    print("Documentul final (Cap 3 si 4) a fost generat in AgriConnect_Diploma_Final.docx")

if __name__ == '__main__':
    main()
