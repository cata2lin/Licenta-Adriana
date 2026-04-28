import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name="Times New Roman", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_font(run, "Times New Roman", 16 if level == 1 else 14)
    run.bold = True
    return h

def main():
    # Vom crea documentul de la zero (sau poți încărca un template dacă adaugi path-ul)
    output_file = 'AgriConnect_Diploma_v1.docx'
    
    doc = Document()
    
    # 1. Pagina de Titlu (simplificată)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Universitatea ...\nFacultatea ...\n\n\n\n")
    set_font(run, size=14)
    
    run = title.add_run("PROIECT DE DIPLOMĂ\n")
    set_font(run, size=24)
    run.bold = True
    
    run = title.add_run("\nDesignul și Implementarea unei Arhitecturi Software Distribuite pentru un Marketplace B2B Agricol\n\n\n")
    set_font(run, size=18)
    run.bold = True
    
    run = title.add_run("Absolvent: ...\nCoordonator: ...\n\n\n2026")
    set_font(run, size=14)
    
    doc.add_page_break()

    # 2. Cuprins
    add_heading(doc, "CUPRINS", 1)
    p = doc.add_paragraph()
    run = p.add_run("[Aici se va genera automat cuprinsul în Microsoft Word]")
    set_font(run, size=12)
    run.italic = True
    doc.add_page_break()

    # 3. Lista de Acronime
    add_heading(doc, "LISTA DE ACRONIME", 1)
    acronyms = [
        ("API", "Application Programming Interface"),
        ("B2B", "Business-to-Business"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("DAC7", "Directiva (UE) 2021/514 privind cooperarea administrativă în domeniul fiscal"),
        ("DDD", "Domain-Driven Design"),
        ("ESBuild", "Extremely Fast JavaScript Bundler"),
        ("GDPR", "General Data Protection Regulation"),
        ("GIN", "Generalized Inverted Index (PostgreSQL)"),
        ("HMR", "Hot Module Replacement"),
        ("IAM", "Identity and Access Management"),
        ("JSONB", "JavaScript Object Notation - Binary"),
        ("JWT", "JSON Web Token"),
        ("KYC", "Know Your Customer"),
        ("ORM", "Object-Relational Mapping"),
        ("OSRM", "Open Source Routing Machine"),
        ("PDF", "Portable Document Format"),
        ("PSD2", "Payment Services Directive 2"),
        ("SPA", "Single Page Application"),
        ("SQL", "Structured Query Language"),
        ("UI", "User Interface"),
        ("UIT", "Codul Unic de Identificare a Transportului (RO e-Transport)"),
        ("UX", "User Experience")
    ]
    
    for acr, desc in acronyms:
        p = doc.add_paragraph()
        run1 = p.add_run(f"{acr}")
        run1.bold = True
        set_font(run1, size=12)
        run2 = p.add_run(f" – {desc}")
        set_font(run2, size=12)
        
    doc.add_page_break()

    # 4. Lista de Figuri
    add_heading(doc, "LISTA DE FIGURI ȘI TABELE", 1)
    p = doc.add_paragraph()
    run = p.add_run("[Aici se va genera automat lista figurilor]")
    set_font(run, size=12)
    run.italic = True
    doc.add_page_break()

    doc.save(output_file)
    print(f"[OK] Front matter generat in {output_file}")

if __name__ == "__main__":
    main()
